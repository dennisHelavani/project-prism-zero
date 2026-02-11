"""
DOCX document generator with placeholder replacement.
Handles Word's run-splitting of placeholders and optional image insertion.
"""

import os
import re
import logging
from typing import Dict, Optional
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class TemplateNotFoundError(Exception):
    """Raised when template file is not found."""
    pass


def generate_docx(
    template_path: str,
    output_path: str,
    placeholders: Dict[str, str],
    images: Optional[Dict[str, str]] = None,
    blue_flags: Optional[Dict[str, bool]] = None
) -> str:
    """
    Generate a DOCX document by replacing placeholders in a template.
    
    Args:
        template_path: Path to the DOCX template file
        output_path: Path where the generated DOCX will be saved
        placeholders: Dictionary of placeholder keys and their replacement values
        images: Optional dictionary of image placeholder keys and image file paths
        blue_flags: Optional dictionary of blue logic flags for conditional section removal
    
    Returns:
        Path to the generated DOCX file
    
    Raises:
        TemplateNotFoundError: If template file doesn't exist
    """
    if not os.path.exists(template_path):
        raise TemplateNotFoundError(f"Template not found: {template_path}")
    
    logger.info(f"Loading template: {template_path}")
    doc = Document(template_path)
    
    images = images or {}
    blue_flags = blue_flags or {}
    logger.info(f"Images dictionary: {images}")
    logger.info(f"Blue flags: {blue_flags}")
    
    # Apply blue logic - remove conditional sections/bullets based on flags
    if blue_flags:
        _apply_blue_logic(doc, blue_flags)
    
    # Process all parts of the document - log paragraphs with placeholders
    for p in doc.paragraphs:
        if '{{' in p.text and '}}' in p.text:
            logger.info(f"Found placeholder in body paragraph: {p.text[:80]}...")
    _replace_placeholders_in_paragraphs(doc.paragraphs, placeholders, images, doc)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_placeholders_in_paragraphs(cell.paragraphs, placeholders, images, doc)
    
    # Process headers and footers
    for section in doc.sections:
        # Header
        if section.header:
            logger.info(f"Processing header with {len(section.header.paragraphs)} paragraphs")
            for p in section.header.paragraphs:
                if p.text.strip():
                    logger.info(f"Header paragraph text: {p.text[:100]}...")
            _replace_placeholders_in_paragraphs(section.header.paragraphs, placeholders, images, doc)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _replace_placeholders_in_paragraphs(cell.paragraphs, placeholders, images, doc)
            # Process text boxes in header
            _replace_placeholders_in_element(section.header._element, placeholders, images, doc)
        
        # Footer
        if section.footer:
            _replace_placeholders_in_paragraphs(section.footer.paragraphs, placeholders, images, doc)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _replace_placeholders_in_paragraphs(cell.paragraphs, placeholders, images, doc)
            # Process text boxes in footer
            _replace_placeholders_in_element(section.footer._element, placeholders, images, doc)
    
    # Process Text Boxes (shapes with text content) in main document body
    # Text boxes are stored in w:txbxContent elements
    _replace_placeholders_in_textboxes(doc, placeholders, images)
    
    # Force Table of Contents to update when document is opened in Word
    _set_update_fields_on_open(doc)
    
    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    
    # Save the document
    logger.info(f"Saving document: {output_path}")
    doc.save(output_path)
    
    return output_path


def _set_update_fields_on_open(doc: Document):
    """
    Set the updateFields property in document settings.
    This forces Word to update all fields (including Table of Contents)
    when the document is opened.
    """
    try:
        from lxml import etree
        
        # Access the document settings part
        settings = doc.settings.element
        
        # Define the namespace
        w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        
        # Check if updateFields already exists
        update_fields = settings.find(f'{w_ns}updateFields')
        
        if update_fields is None:
            # Create the updateFields element
            update_fields = etree.SubElement(settings, f'{w_ns}updateFields')
        
        # Set the val attribute to "true"
        update_fields.set(f'{w_ns}val', 'true')
        
        logger.info("Set updateFields=true - TOC will update when document is opened")
        
    except Exception as e:
        logger.warning(f"Could not set updateFields property: {e}")



def _replace_placeholders_in_element(element, placeholders: Dict[str, str], images: Dict[str, str], doc: Document):
    """
    Process text boxes within a specific XML element (header, footer, etc.).
    """
    from docx.text.paragraph import Paragraph
    
    # Find all txbxContent elements in this element's subtree
    for txbx_content in element.iter(qn('w:txbxContent')):
        for p_elem in txbx_content.iter(qn('w:p')):
            try:
                # Pass doc.part so that add_picture() works
                paragraph = Paragraph(p_elem, doc.part)
                full_text = paragraph.text
                if '{{' in full_text and '}}' in full_text:
                    logger.info(f"Found placeholder in header/footer text box: {full_text[:50]}...")
                    _replace_in_paragraph(paragraph, placeholders, images, doc)
            except Exception as e:
                logger.warning(f"Error processing header/footer text box: {e}")


def _replace_placeholders_in_textboxes(doc: Document, placeholders: Dict[str, str], images: Dict[str, str]):
    """
    Process text boxes (shapes with text content).
    Text boxes in DOCX are stored in w:txbxContent elements within shapes.
    This handles both VML text boxes and DrawingML text boxes.
    """
    from docx.text.paragraph import Paragraph
    
    # Find all txbxContent elements (text box content containers)
    # These can be in VML shapes (w:pict/v:shape/v:textbox/w:txbxContent)
    # or in DrawingML (w:drawing/.../wps:txbx/w:txbxContent)
    
    txbx_contents = doc.element.iter(qn('w:txbxContent'))
    
    for txbx_content in txbx_contents:
        # Find all paragraph elements inside the text box
        for p_elem in txbx_content.iter(qn('w:p')):
            try:
                # Pass doc.part so that add_picture() works
                paragraph = Paragraph(p_elem, doc.part)
                
                # Get the text to check for placeholders
                full_text = paragraph.text
                if '{{' in full_text and '}}' in full_text:
                    logger.info(f"Found placeholder in text box: {full_text[:50]}...")
                    _replace_in_paragraph(paragraph, placeholders, images, doc)
                    
            except Exception as e:
                logger.warning(f"Error processing text box paragraph: {e}")


def _insert_multiline_content(paragraph, content: str, placeholder_key: str) -> bool:
    """
    Insert multi-line content as separate paragraphs after the original paragraph.
    Each line becomes a new paragraph inheriting the original's list style.
    
    Used for AI_CONSTRUCTION_SEQUENCE and AI_RISK_MANAGEMENT which contain
    newline-separated list items.
    
    For AI_RISK_MANAGEMENT, subheaders (lines starting with "7.") are styled 
    with Calibri 12pt orange.
    
    Returns:
        True if content was inserted (original paragraph should be removed)
    """
    if not content or '\n' not in content:
        return False  # Not multi-line, use normal replacement
    
    lines = [line.strip() for line in content.split('\n') if line.strip()]
    if len(lines) <= 1:
        return False  # Only one line, use normal replacement
    
    logger.info(f"Inserting {len(lines)} list items for {placeholder_key}")
    
    # Get the parent element and position
    parent = paragraph._element.getparent()
    if parent is None:
        return False
    
    # Find the index of this paragraph
    para_index = list(parent).index(paragraph._element)
    
    # Get paragraph properties to copy (list style, etc.)
    original_pPr = paragraph._element.find(qn('w:pPr'))
    
    # Insert each line as a new paragraph AFTER the original
    from docx.oxml.ns import nsmap
    from docx.oxml import OxmlElement
    from copy import deepcopy
    from docx.shared import Pt, RGBColor
    import re
    
    # Pattern to detect subheaders like "7.1 WORKING AT HEIGHT", "7.2 LIFTING OPERATIONS"
    subheader_pattern = re.compile(r'^7\.\d+\s+')
    # Pattern to detect and strip fake bullet characters at start of line
    fake_bullet_pattern = re.compile(r'^[•\-\*]\s*')
    
    for i, line in enumerate(lines):
        # Create a new paragraph element
        new_p = OxmlElement('w:p')
        
        # Check if this is an AI_RISK_MANAGEMENT subheader (starts with 7.x)
        is_subheader = placeholder_key == 'AI_RISK_MANAGEMENT' and subheader_pattern.match(line)
        # Control measure = AI_RISK_MANAGEMENT line that is NOT a subheader and NOT empty
        is_control_measure = placeholder_key == 'AI_RISK_MANAGEMENT' and not is_subheader and line.strip()
        
        # Clean up the line text
        clean_line = line
        if is_subheader:
            # Remove trailing colon or period from subheader
            clean_line = line.rstrip(':.')
        elif is_control_measure:
            # Strip any fake bullet characters from the start
            clean_line = fake_bullet_pattern.sub('', line)
        
        # Copy paragraph properties (preserves list style) - but NOT for control measures
        if original_pPr is not None and not is_control_measure:
            new_pPr = deepcopy(original_pPr)
            # Remove any shading from the copied properties
            shd_elements = new_pPr.findall(qn('w:shd'))
            for shd in shd_elements:
                new_pPr.remove(shd)
            new_p.append(new_pPr)
        
        # Create a run with the text
        new_r = OxmlElement('w:r')
        new_t = OxmlElement('w:t')
        new_t.text = clean_line
        new_r.append(new_t)
        
        if is_subheader:
            # Apply Calibri 12pt Orange styling to the run
            rPr = OxmlElement('w:rPr')
            
            # Font: Calibri
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Calibri')
            rFonts.set(qn('w:hAnsi'), 'Calibri')
            rFonts.set(qn('w:cs'), 'Calibri')
            rPr.append(rFonts)
            
            # Size: 12pt (24 half-points)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '24')  # 12pt = 24 half-points
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '24')
            rPr.append(szCs)
            
            # Color: Orange (F65C02)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), 'F65C02')  # Orange hex color (RGB: 246, 92, 2)
            rPr.append(color)
            
            # Bold for subheaders
            bold = OxmlElement('w:b')
            rPr.append(bold)
            
            # Insert rPr at the beginning of the run
            new_r.insert(0, rPr)
            
            # Add paragraph spacing: 6pt before, single line spacing
            # Get or create paragraph properties
            pPr = new_p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                new_p.insert(0, pPr)
            
            # Add spacing element: 6pt before (120 twentieths of a point), single line spacing
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '120')  # 6pt = 120 twentieths of a point
            spacing.set(qn('w:line'), '240')    # Single line spacing (240 twentieths = 12pt)
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)
            
            logger.debug(f"Applied subheader styling to: {clean_line[:40]}...")
        
        elif is_control_measure:
            # Apply 10pt font size to control measure lines
            rPr = OxmlElement('w:rPr')
            
            # Size: 10pt (20 half-points)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '20')  # 10pt = 20 half-points
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '20')
            rPr.append(szCs)
            
            # Insert rPr at the beginning of the run
            new_r.insert(0, rPr)
            
            # Add bullet formatting to paragraph
            pPr = OxmlElement('w:pPr')
            
            # Create numbering properties for bullet
            numPr = OxmlElement('w:numPr')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')  # Level 0 (top level)
            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), '1')  # Use numbering definition 1 (usually bullets)
            numPr.append(ilvl)
            numPr.append(numId)
            pPr.append(numPr)
            
            # Add left indent for bullet (0.25 inch = 360 twips for hanging, 0.5 inch = 720 twips for left)
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '720')      # 0.5 inch left margin
            ind.set(qn('w:hanging'), '360')   # 0.25 inch hanging indent
            pPr.append(ind)
            
            # Add run properties to paragraph to style the bullet symbol at 10pt
            pPr_rPr = OxmlElement('w:rPr')
            pPr_sz = OxmlElement('w:sz')
            pPr_sz.set(qn('w:val'), '20')  # 10pt = 20 half-points
            pPr_rPr.append(pPr_sz)
            pPr_szCs = OxmlElement('w:szCs')
            pPr_szCs.set(qn('w:val'), '20')
            pPr_rPr.append(pPr_szCs)
            pPr.append(pPr_rPr)
            
            new_p.insert(0, pPr)
        
        new_p.append(new_r)
        
        # Insert after the original paragraph (in reverse order so they end up in correct order)
        parent.insert(para_index + 1 + i, new_p)
    
    # Remove the original placeholder paragraph
    parent.remove(paragraph._element)
    
    return True


def _insert_risk_assessment_table(paragraph, content: str, doc: Document) -> bool:
    """
    Insert a risk assessment table from pipe-separated content.
    
    Format expected:
    Activity|Hazard|Persons at Risk|Control Measures|Residual Risk
    Row1Col1|Row1Col2|Row1Col3|Row1Col4|Row1Col5
    Row2Col1|Row2Col2|Row2Col3|Row2Col4|Row2Col5
    
    Control Measures may contain semicolon-separated bullet points.
    
    Returns:
        True if table was inserted successfully
    """
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    if not content or '|' not in content:
        logger.warning("AI_RISK_ASSESSMENT content doesn't contain pipe separators")
        return False
    
    # Parse rows from content
    rows_data = [line.strip() for line in content.split('\n') if line.strip() and '|' in line]
    
    if len(rows_data) < 1:
        logger.warning("AI_RISK_ASSESSMENT has no valid rows")
        return False
    
    logger.info(f"Creating risk assessment table with {len(rows_data)} rows")
    
    # Get the parent element to insert table
    parent = paragraph._element.getparent()
    if parent is None:
        return False
    
    # Find paragraph position
    para_index = list(parent).index(paragraph._element)
    
    # Define headers
    headers = ["Activity", "Hazard", "Persons at Risk", "Control Measures", "Residual Risk"]
    num_cols = 5
    
    # Create table with header row + data rows
    table = doc.add_table(rows=len(rows_data) + 1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style the table
    try:
        table.style = 'Table Grid'
    except:
        pass  # Style may not exist in template
    
    # Add header row
    header_row = table.rows[0]
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header_text
        # Bold and style header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header background color (dark blue)
        try:
            from docx.oxml import OxmlElement
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '1F4E79')  # Dark blue
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
            # White text for header
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
        except Exception as e:
            logger.debug(f"Could not set header styling: {e}")
    
    # Add data rows
    for row_idx, row_data in enumerate(rows_data):
        cols = [col.strip() for col in row_data.split('|')]
        data_row = table.rows[row_idx + 1]
        
        for col_idx in range(min(len(cols), num_cols)):
            cell = data_row.cells[col_idx]
            cell_text = cols[col_idx] if col_idx < len(cols) else ""
            
            # Control Measures column (index 3) - convert semicolons to bullet points
            if col_idx == 3 and ';' in cell_text:
                # Clear existing content
                cell.text = ""
                items = [item.strip() for item in cell_text.split(';') if item.strip()]
                for item_idx, item in enumerate(items):
                    if item_idx == 0:
                        cell.paragraphs[0].text = f"• {item}"
                    else:
                        cell.add_paragraph(f"• {item}")
                # Style all paragraphs
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
            else:
                cell.text = cell_text
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
        
        # Alternate row shading (light gray for even rows)
        if row_idx % 2 == 1:
            try:
                for cell in data_row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'F2F2F2')  # Light gray
                    shd.set(qn('w:val'), 'clear')
                    tcPr.append(shd)
            except:
                pass
    
    # Set column widths
    try:
        table.columns[0].width = Inches(1.2)  # Activity
        table.columns[1].width = Inches(1.2)  # Hazard
        table.columns[2].width = Inches(1.0)  # Persons at Risk
        table.columns[3].width = Inches(2.5)  # Control Measures (wider)
        table.columns[4].width = Inches(0.8)  # Residual Risk
    except:
        pass
    
    # Move table to paragraph position
    try:
        tbl_element = table._tbl
        parent.insert(para_index, tbl_element)
        # Remove the original placeholder paragraph
        parent.remove(paragraph._element)
        logger.info("Risk assessment table inserted successfully")
        return True
    except Exception as e:
        logger.error(f"Failed to insert risk assessment table: {e}")
        return False


def _replace_placeholders_in_paragraphs(
    paragraphs,
    placeholders: Dict[str, str],
    images: Dict[str, str],
    doc: Document
):
    """
    Replace placeholders in a list of paragraphs.
    Handles Word's run-splitting of placeholders.
    Deletes paragraphs that become empty after placeholder replacement.
    """
    # Create a list of paragraphs to remove
    paragraphs_to_remove = []
    
    # AI placeholders that should be split into multiple lines
    MULTILINE_PLACEHOLDERS = ['AI_WORK_LIST', 'AI_CONSTRUCTION_SEQUENCE', 'AI_RISK_MANAGEMENT', 'AI_SEQUENCE_OF_WORKS']
    
    # Make a copy of the list since we may modify it during iteration
    for paragraph in list(paragraphs):
        para_text = paragraph.text
        
        # Special handling for AI_RISK_ASSESSMENT - convert to table
        if '{{AI_RISK_ASSESSMENT}}' in para_text and 'AI_RISK_ASSESSMENT' in placeholders:
            content = placeholders['AI_RISK_ASSESSMENT']
            if content and '|' in content:
                if _insert_risk_assessment_table(paragraph, content, doc):
                    continue  # Table was inserted, skip other processing
        
        # Check if this paragraph contains a multi-line AI placeholder
        multiline_handled = False
        for placeholder_key in MULTILINE_PLACEHOLDERS:
            full_placeholder = '{{' + placeholder_key + '}}'
            if full_placeholder in para_text and placeholder_key in placeholders:
                content = placeholders[placeholder_key]
                if content and '\n' in content:
                    # Use the multi-line helper
                    if _insert_multiline_content(paragraph, content, placeholder_key):
                        multiline_handled = True
                        break
        
        if multiline_handled:
            continue  # Skip normal replacement, paragraph was already handled
        
        # Normal replacement for other placeholders
        was_replaced_with_empty = _replace_in_paragraph(paragraph, placeholders, images, doc)
        
        # If a placeholder was replaced with empty string, check if paragraph is now empty
        if was_replaced_with_empty:
            # Check if the text is empty or just whitespace
            if not paragraph.text.strip():
                paragraphs_to_remove.append(paragraph)
    
    # Remove the empty paragraphs
    for p in paragraphs_to_remove:
        try:
            p._element.getparent().remove(p._element)
            logger.debug("Removed empty paragraph after placeholder replacement")
        except Exception as e:
            logger.warning(f"Failed to remove empty paragraph: {e}")


def _replace_in_paragraph(
    paragraph,
    placeholders: Dict[str, str],
    images: Dict[str, str],
    doc: Document
) -> bool:
    """
    Replace placeholders in a single paragraph.
    Uses a two-pass approach:
    1. First, merge runs that contain split placeholders
    2. Then, perform the actual replacement
    
    Returns:
        bool: True if a placeholder was replaced with an empty string (candidate for deletion)
    """
    # Get full paragraph text
    full_text = paragraph.text
    
    # Find all placeholder patterns {{KEY}}
    placeholder_pattern = r'\{\{([A-Za-z0-9_]+)\}\}'
    matches = list(re.finditer(placeholder_pattern, full_text))
    
    if not matches:
        return False  # No placeholders in this paragraph
    
    # Check if placeholders are split across runs
    # If so, we need to merge the runs first
    if len(paragraph.runs) > 1:
        _merge_placeholder_runs(paragraph)
        
    replaced_with_empty = False
    
    # Now perform replacements on each run
    for run in paragraph.runs:
        run_text = run.text
        
        # Find placeholders in this run
        for match in re.finditer(placeholder_pattern, run_text):
            placeholder_key = match.group(1)
            full_placeholder = match.group(0)  # {{KEY}}
            
            # Check if this is an image placeholder
            if placeholder_key in images and images[placeholder_key]:
                image_path = images[placeholder_key]
                if os.path.exists(image_path):
                    # Replace with image
                    logger.info(f"Inserting image for {placeholder_key}: {image_path}")
                    run.text = run_text.replace(full_placeholder, "")
                    run_text = run.text
                    try:
                        # Special sizing for different logo placeholders
                        if placeholder_key == "CPP_LOGO_TOP_RIGHT_IMG":
                            # 0.51 inch x 0.51 inch for header logo
                            run.add_picture(image_path, width=Inches(0.51), height=Inches(0.51))
                        elif placeholder_key == "CPP_LOGO_COVER_MIDDLE_IMG":
                            # 1.43 inch x 1.43 inch for cover page logo
                            run.add_picture(image_path, width=Inches(1.43), height=Inches(1.43))
                        elif placeholder_key == "RAMS_COVER_PAGE_LOGO_IMG":
                            # 1.41 inch x 1.41 inch for RAMS cover page logo
                            run.add_picture(image_path, width=Inches(1.41), height=Inches(1.41))
                        elif placeholder_key == "RAMS_CLIENT_PAGE_LOGO_IMG":
                            # 1.41 inch x 1.41 inch for RAMS client logo
                            run.add_picture(image_path, width=Inches(1.41), height=Inches(1.41))
                        elif placeholder_key == "RAMS_DELIVERIES_IMG":
                            # 6.26 inch width x 4.18 inch height for RAMS deliveries TMP image
                            run.add_picture(image_path, width=Inches(6.26), height=Inches(4.18))
                        elif placeholder_key == "RAMS_FIRE_PLAN_IMG":
                            # 6.26 inch width x 4.18 inch height for RAMS fire plan image
                            run.add_picture(image_path, width=Inches(6.26), height=Inches(4.18))
                        elif placeholder_key == "RAMS_NEAREST_HOSPITAL_IMG":
                            # 2.09 inch width x 3.13 inch height for RAMS hospital route map
                            run.add_picture(image_path, width=Inches(2.09), height=Inches(3.13))
                        else:
                            # Default 2.0 inches width for other images
                            run.add_picture(image_path, width=Inches(2.0))
                    except Exception as e:
                        logger.warning(f"Failed to insert image {image_path}: {e}")
                else:
                    # Image path doesn't exist, remove placeholder
                    logger.warning(f"Image not found for {placeholder_key}: {image_path}")
                    run.text = run_text.replace(full_placeholder, "")
                    run_text = run.text
                    replaced_with_empty = True
            
            # Check if this is a text placeholder
            elif placeholder_key in placeholders:
                val = placeholders[placeholder_key]
                if not val and placeholder_key in ["RAMS_COVER_PAGE_LOGO_IMG", "RAMS_CLIENT_PAGE_LOGO_IMG"]:
                    # Specific request: 3 empty lines if logo missing
                    replacement = "\n\n\n"
                elif not val and placeholder_key.startswith("RAMS_PERMIT_"):
                    # Specific request: Remove the line completely for empty permits
                    # We set replacement to empty string, and ensure we DON'T use NBSP
                    replacement = ""
                elif not val and placeholder_key.startswith("include_"):
                    # Blue Flag placeholders: Use empty string so the paragraph gets removed
                    replacement = ""
                    replaced_with_empty = True
                else:
                    replacement = str(val) if val else "\u00A0"  # Non-breaking space
                
                if not val and not placeholder_key.startswith("RAMS_PERMIT_") and not placeholder_key.startswith("include_"):
                    # Do NOT mark as empty if we used a non-breaking space (for normal fields)
                    # This prevents the paragraph from being removed, preserving the table cell
                    pass 
                elif not val and placeholder_key.startswith("RAMS_PERMIT_"):
                    # DO mark as empty for permits, so the paragraph gets removed
                    replaced_with_empty = True 
                
                # SPECIAL HANDLING: Titles should be vertical (bottom-to-top)
                if placeholder_key in ["RAMS_TITLE", "CPP_PROJECT_TITLE", "CPP_TITLE"]:
                    # If this is CPP_TITLE or RAMS_TITLE in the specific text box context
                    if placeholder_key in ["CPP_TITLE", "RAMS_TITLE"]:
                        # Ensure it starts on a new line
                        if not str(replacement).startswith('\n'):
                            replacement = '\n' + str(replacement)
                        
                        # Apply Calibri font
                        run.font.name = 'Calibri'
                        from docx.shared import Pt
                        # Optional: set size if needed, but inheriting is usually safer unless requested
                        # run.font.size = Pt(11) 
                        
                    _set_cell_text_direction_vertical(paragraph)
                
                # SPECIAL HANDLING: Timestamps (Roboto Mono, 10pt)
                if placeholder_key in ["RAMS_DATE_TIME", "RAMS_DATE_STAMPED"]:
                    run.font.name = 'Roboto Mono'
                    from docx.shared import Pt
                    run.font.size = Pt(10)

                run.text = run_text.replace(full_placeholder, replacement)
                run_text = run.text
                
                # Clear any background highlighting/shading (template placeholders may have colored backgrounds)
                # 1. Clear run highlight color
                run.font.highlight_color = None
                
                # 2. Clear run shading (background color in the run's XML)
                try:
                    rPr = run._element.get_or_add_rPr()
                    shd_elements = rPr.findall(qn('w:shd'))
                    for shd in shd_elements:
                        rPr.remove(shd)
                except Exception as e:
                    logger.debug(f"Could not remove run shading: {e}")
                
                # 3. Clear paragraph shading if this is a placeholder paragraph
                try:
                    pPr = paragraph._element.get_or_add_pPr()
                    shd_elements = pPr.findall(qn('w:shd'))
                    for shd in shd_elements:
                        pPr.remove(shd)
                except Exception as e:
                    logger.debug(f"Could not remove paragraph shading: {e}")
            
            # Placeholder not in data - use non-breaking space to preserve cell structure
            else:
                logger.debug(f"Placeholder not found in data: {placeholder_key}")
                run.text = run_text.replace(full_placeholder, "\u00A0")  # Non-breaking space
                run_text = run.text
                # Do NOT mark as empty if we used a non-breaking space
                # replaced_with_empty = True 
                
    return replaced_with_empty


def _set_cell_text_direction_vertical(paragraph):
    """
    Set the text direction of the parent table cell to 'btLr' (bottom-to-top).
    Also sets vertical alignment to 'center' and paragraph alignment to 'left' (bottom anchor).
    Requires the paragraph to be inside a table cell.
    """
    try:
        # 1. Force Paragraph Alignment to LEFT (which is Bottom for vertical text)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Access the underlying OpenXML element
        p_element = paragraph._element
        
        # Traverse up to find the Table Cell (tc) element
        current = p_element
        while current is not None:
            if current.tag.endswith('tc'): # w:tc
                 tc = current
                 tcPr = tc.get_or_add_tcPr()
                 
                 # 2. Set Text Direction: btLr (bottom to top, left to right)
                 # This rotates text 90deg CCW. Bottom of letter faces East.
                 text_direction = qn('w:textDirection')
                 matches = tcPr.xpath('w:textDirection')
                 if matches:
                     matches[0].set(qn('w:val'), 'btLr')
                 else:
                     elem = OxmlElement('w:textDirection')
                     elem.set(qn('w:val'), 'btLr')
                     tcPr.append(elem)
                 
                 # 3. Set Vertical Alignment: Center
                 v_align = qn('w:vAlign')
                 matches_va = tcPr.xpath('w:vAlign')
                 if matches_va:
                     matches_va[0].set(qn('w:val'), 'center')
                 else:
                     elem_va = OxmlElement('w:vAlign')
                     elem_va.set(qn('w:val'), 'center')
                     tcPr.append(elem_va)
                     
                 logger.info("Applied vertical text direction (btLr) + center align to table cell")
                 return
            current = current.getparent()
            
        logger.warning("Could not find parent Table Cell for vertical title - make sure it's in a table!")
            
    except Exception as e:
        logger.warning(f"Failed to set vertical text direction: {e}")


def OxmlElement(tag):
    """Helper to create OXML element"""
    from docx.oxml.ns import nsmap
    from lxml import etree
    return etree.Element(qn(tag), nsmap=nsmap)


def _merge_placeholder_runs(paragraph):
    """
    Merge runs that contain split placeholders.

    
    Word sometimes splits text like "{{KEY}}" into multiple runs:
    Run 1: "{{"
    Run 2: "KEY"
    Run 3: "}}"
    
    This function detects and merges such split placeholders.
    """
    if not paragraph.runs:
        return
    
    # Get full paragraph text and check for placeholders
    full_text = paragraph.text
    if "{{" not in full_text or "}}" not in full_text:
        return
    
    # Find positions of {{ and }} in the text
    open_pos = []
    close_pos = []
    
    i = 0
    while i < len(full_text) - 1:
        if full_text[i:i+2] == "{{":
            open_pos.append(i)
            i += 2
        elif full_text[i:i+2] == "}}":
            close_pos.append(i)
            i += 2
        else:
            i += 1
    
    if not open_pos or not close_pos:
        return
    
    # Build a map of character positions to runs
    char_to_run = []
    for run_idx, run in enumerate(paragraph.runs):
        for _ in range(len(run.text)):
            char_to_run.append(run_idx)
    
    # Check each placeholder span
    for op in open_pos:
        # Find matching close
        matching_close = None
        for cp in close_pos:
            if cp > op:
                matching_close = cp
                break
        
        if matching_close is None:
            continue
        
        # Check if placeholder spans multiple runs
        if len(char_to_run) <= matching_close + 1:
            continue
            
        start_run = char_to_run[op] if op < len(char_to_run) else -1
        end_run = char_to_run[matching_close + 1] if matching_close + 1 < len(char_to_run) else char_to_run[-1]
        
        if start_run != end_run and start_run >= 0:
            # Placeholder is split - merge all text into first run
            # This is a simplified approach: merge all runs into one
            merged_text = full_text
            
            # Clear all runs except first
            first_run = paragraph.runs[0]
            first_run.text = merged_text
            
            # Remove other runs (by clearing their text)
            for run in paragraph.runs[1:]:
                run.text = ""
            
            # Only need to do this once per paragraph
            return


# ============================================================================
# Blue Logic: Conditional Section/Bullet Removal
# ============================================================================

# Mapping of blue flags to the text patterns they control
BLUE_FLAG_TEXT_PATTERNS = {
    # Section 3.2 - Competence bullets
    "BLUE_FLAG_COMPETENCE_PLANT_HEIGHT": [
        "CSCS",  # "All personnel to hold CSCS... CPCS/NPORS/IPAF"
        "CPCS",
        "NPORS",
        "IPAF",
        "CPCS/NPORS/IPAF",
        "Site Induction mandatory",  # "Site Induction mandatory; include height, lifting"
    ],
    
    # Section 4.1 & 15 - Heavy logistics/laydown
    "BLUE_FLAG_HEAVY_LOGISTICS_LAYDOWN": [
        "Laydown plans",
        "exclusion zones for loading",
        "laydown areas",
        "crane base",
        "crane pad",
        "hoist(s)",
        "fuel store",
    ],
    
    # Section 4.2 - Public traffic management
    "BLUE_FLAG_PUBLIC_TRAFFIC_MGMT": [
        "Delivery booking system",
        "off-peak deliveries",
        "off‑peak deliveries",  # Different dash character
        "TTRO",
        "permits for footway",
        "footway closures",
    ],
    
    # Section 4.3 - Secure perimeter/hoarding
    "BLUE_FLAG_SECURE_PERIMETER_HOARDING": [
        "Hoarding to suitable standard",
        "secure gates",
        "CCTV",
        "intruder alarms",
    ],
}

# Section numbers to remove entirely - patterns must be very specific to avoid false matches
# These should match ONLY the section headers, not content paragraphs that happen to contain numbers
BLUE_FLAG_SECTION_HEADERS = {
    # Section 16 header patterns - must be at start of line with proper formatting
    "BLUE_FLAG_SECTION_16_FACADE": [
        "16 Access",
        "16. Access",
        "16  Access",
        "16 ACCESS",
        "16. ACCESS",
        "Scaffolding & Façade",
        "Scaffolding & Facade",
        "Access, Scaffolding",
    ],
    # Section 17 header patterns - must be at start of line with proper formatting
    "BLUE_FLAG_SECTION_17_COMMISSIONING": [
        "17 Testing",
        "17. Testing",
        "17  Testing",
        "17 TESTING",
        "17. TESTING",
        "Testing, Commissioning",
        "Commissioning & Handover",
    ],
}


def _apply_blue_logic(doc, blue_flags: Dict[str, bool]):
    """
    Apply blue logic to conditionally remove sections and bullets from the document.
    
    Args:
        doc: The python-docx Document object
        blue_flags: Dictionary of flag names to boolean values (True = keep, False = remove)
    """
    logger.info("Applying blue logic for conditional content removal...")
    
    # Track what we removed for logging
    removed_count = 0
    
    # Process paragraphs for bullet/text pattern removal
    for flag_name, patterns in BLUE_FLAG_TEXT_PATTERNS.items():
        if not blue_flags.get(flag_name, True):  # If flag is False, remove matching content
            removed_count += _remove_paragraphs_matching_patterns(doc, patterns, flag_name)
    
    # Process entire sections for removal (Sections 16 & 17)
    for flag_name, header_patterns in BLUE_FLAG_SECTION_HEADERS.items():
        if not blue_flags.get(flag_name, True):  # If flag is False, remove the section
            removed_count += _remove_section_by_header(doc, header_patterns, flag_name)
    
    logger.info(f"Blue logic applied: removed {removed_count} elements")


def _remove_paragraphs_matching_patterns(doc, patterns: list, flag_name: str) -> int:
    """
    Remove paragraphs (bullets) that contain any of the specified patterns.
    
    Returns:
        Number of paragraphs removed
    """
    removed = 0
    paragraphs_to_remove = []
    
    # Check body paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        for pattern in patterns:
            if pattern.lower() in text.lower():
                paragraphs_to_remove.append(p)
                logger.debug(f"Marking for removal ({flag_name}): {text[:50]}...")
                break
    
    # Check table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    for pattern in patterns:
                        if pattern.lower() in text.lower():
                            paragraphs_to_remove.append(p)
                            logger.debug(f"Marking for removal in table ({flag_name}): {text[:50]}...")
                            break
    
    # Remove the marked paragraphs
    for p in paragraphs_to_remove:
        try:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
                removed += 1
        except Exception as e:
            logger.warning(f"Failed to remove paragraph: {e}")
    
    if removed > 0:
        logger.info(f"Removed {removed} paragraphs for {flag_name}=False")
    
    return removed


def _remove_section_by_header(doc, header_patterns: list, flag_name: str) -> int:
    """
    Remove an entire section starting from a header that matches the patterns,
    until the next Heading 1 style paragraph is encountered.
    
    Returns:
        Number of elements removed
    """
    removed = 0
    in_section = False
    elements_to_remove = []
    
    # Known section headers that should stop removal (sections 18, 19, 20+)
    stop_patterns = [
        "Programme",
        "Records & Documentation",
        "Records \u0026 Documentation",
        "Appendices",
        "Sign-Off",
        "Sign‑Off",
    ]
    
    for p in doc.paragraphs:
        text = p.text.strip()
        style_name = p.style.name if p.style else ""
        
        # Check if this is the start of the section to remove
        if not in_section:
            for pattern in header_patterns:
                if pattern in text:
                    in_section = True
                    elements_to_remove.append(p)
                    logger.debug(f"Started section removal at: {text[:50]}...")
                    break
        else:
            # Check if we've hit: 
            # 1. A Heading 1 style paragraph (new major section)
            # 2. A known section header that should stop removal
            # 3. A numbered section header like "18." or "19."
            is_next_section = False
            
            # Check for Heading 1 style
            if "Heading 1" in style_name and not any(pat in text for pat in header_patterns):
                is_next_section = True
            
            # Check for known stop patterns
            if any(stop_pat in text for stop_pat in stop_patterns):
                is_next_section = True
            
            # Check for numbered section headers
            if re.match(r'^\d{1,2}\.?\s', text) and not any(pat in text for pat in header_patterns):
                is_next_section = True
            
            if is_next_section:
                in_section = False
                logger.debug(f"Ended section removal at: {text[:50]}...")
            else:
                elements_to_remove.append(p)
    
    # Remove the elements
    for p in elements_to_remove:
        try:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
                removed += 1
        except Exception as e:
            logger.warning(f"Failed to remove section element: {e}")
    
    if removed > 0:
        logger.info(f"Removed {removed} elements for section {flag_name}=False")
    
    return removed
